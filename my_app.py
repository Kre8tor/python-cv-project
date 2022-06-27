from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(txt):
    pyttsx3.speak(txt)

document = Document()

# profile picture
document.add_picture('Nicol.png',width=Inches(2.0))

# name phone number and e-mail details
name = input('What is your name? ')
speak('Hello' + name + ' How are you today ?')

speak('What is your phone number? ')
phone_number = input('What is your phone number? ')
speak(name + 'What is your e-mail address')
email = input('What is your e-mail address? ')

document.add_paragraph(
    name + ' | ' +  phone_number  + ' | ' +  email )

# about me 
document.add_heading('About me')
document.add_paragraph(input('Tell me about yourself? '))

# Work experience 
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From Date ')
to_date = input('To Date ')


p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company)

p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No ' )
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
                'Describe your experience at ' + company + ' ')
        p.add_run(experience_details)
    else:
        break
# Skills  
document.add_heading('Skills list')
skill = input('Please list your skills : ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills ? yes or no ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter Skill : ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using python 3.9 by Nicol Posthumus"

    
document.save('cv.docx')

